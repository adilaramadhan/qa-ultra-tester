#!/usr/bin/env python3
"""
QA Ultra Tester - Report Generator v2.1
Generates premium HTML and Word (.docx) reports from test-results.json
Single source of truth: hasil-test/test-results.json

Changes v2.1:
- Fixed dark/light toggle
- Fixed double numbering in Steps to Reproduce
- Added video embed (test recording)
- Added screenshot on failed/bug test results
- Performance metrics now show full names + descriptions
- Duration converted from ms to human-readable (min:sec)
"""

import os
import sys
import json
import base64
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ── Performance metric descriptions ──────────────────────────
METRIC_INFO = {
    "LCP": {"full": "Largest Contentful Paint", "desc": "Waktu render elemen konten utama/terbesar (gambar/banner/teks) di viewport. Target standar: < 2.5 detik (Good), 2.5s - 4.0s (Needs Improvement), > 4.0s (Poor)."},
    "FID": {"full": "First Input Delay", "desc": "Waktu respon dari klik/interaksi pertama pengguna hingga browser mulai memproses event. Target standar: < 100ms (Good), 100ms - 300ms (Needs Improvement), > 300ms (Poor)."},
    "CLS": {"full": "Cumulative Layout Shift", "desc": "Skor stabilitas pergeseran elemen visual yang tidak terduga saat halaman sedang dimuat. Target standar: < 0.1 (Good), 0.1 - 0.25 (Needs Improvement), > 0.25 (Poor)."},
    "TTFB": {"full": "Time to First Byte", "desc": "Durasi waktu dari pengiriman request HTTP sampai byte data pertama diterima dari server/backend. Target standar: < 800ms (Good), > 1800ms (Poor)."},
    "FCP": {"full": "First Contentful Paint", "desc": "Waktu saat elemen konten pertama (teks, gambar latar, canvas) mulai terlihat di layar pengguna. Target standar: < 1.8 detik (Good), 1.8s - 3.0s (Needs Improvement)."},
    "INP": {"full": "Interaction to Next Paint", "desc": "Mengukur latensi keseluruhan respon UI terhadap semua interaksi pengguna (klik, ketik, tap) di halaman. Target standar: < 200ms (Good)."},
    "TBT": {"full": "Total Blocking Time", "desc": "Total durasi waktu main thread browser terhalang/freeze antara FCP dan Time to Interactive. Target standar: < 200ms (Good)."},
    "SI": {"full": "Speed Index", "desc": "Indeks kecepatan visual seberapa cepat seluruh konten di atas lipatan layar tampil lengkap. Target standar: < 3.4 detik (Good)."},
    "TTI": {"full": "Time to Interactive", "desc": "Waktu yang dibutuhkan hingga halaman benar-benar siap dan responsif menerima input pengguna. Target standar: < 3.8 detik (Good)."},
    "PAGE LOAD": {"full": "Full Page Load Time", "desc": "Total waktu dari navigasi awal hingga seluruh resource, script, dan asset halaman selesai dimuat sepenuhnya. Target standar: < 3.0 detik."},
    "AVG RESPONSE TIME": {"full": "Average Response Time", "desc": "Rata-rata waktu pemrosesan permintaan API/layanan backend dari client hingga selesai. Target standar: < 500ms."},
    "P95 LATENCY": {"full": "95th Percentile Latency", "desc": "Batas latensi di mana 95% dari seluruh request selesai di bawah waktu ini. Target standar: < 1000ms."},
    "THROUGHPUT": {"full": "API Throughput / Request per Second", "desc": "Jumlah request per detik (RPS) yang berhasil dilayani secara concurrent tanpa kegagalan."},
}


def format_duration(ms):
    """Convert milliseconds to human-readable format."""
    if ms is None or ms == 0:
        return "0s"
    total_seconds = ms / 1000
    if total_seconds < 1:
        return f"{ms}ms"
    minutes = int(total_seconds // 60)
    seconds = total_seconds % 60
    if minutes == 0:
        return f"{seconds:.1f}s"
    return f"{minutes}m {seconds:.1f}s"


def strip_step_number(step_text):
    """Remove leading number prefix like '1. ' or '1) ' from step text."""
    import re
    return re.sub(r'^\d+[\.\)\-]\s*', '', str(step_text))


class QAReportGenerator:
    def __init__(self, output_dir="hasil-test"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.data = {}
        # Auto-load test-results.json if present
        default_json = self.output_dir / "test-results.json"
        if default_json.exists():
            try:
                with open(default_json, 'r', encoding='utf-8-sig') as f:
                    self.data = json.load(f)
            except Exception as e:
                pass

    def load_from_json(self, filename="test-results.json"):
        """Load all data from the single source of truth JSON file."""
        filepath = self.output_dir / filename
        if not filepath.exists():
            print(f"Error: {filepath} not found. Run tests first.")
            return False
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            self.data = json.load(f)
        return True

    def set_data(self, data: dict):
        """Set data directly (for programmatic use)."""
        self.data = data

    # ── helpers ──────────────────────────────────────────────
    def _bugs(self):
        return self.data.get("bugs", [])

    def _tests(self):
        return self.data.get("test_results", [])

    def _perf(self):
        return self.data.get("performance_metrics", self.data.get("performance", {}))

    def _a11y(self):
        return self.data.get("accessibility", [])

    def _summary(self):
        return self.data.get("summary", {})

    def _quality_score(self):
        explicit_score = self.data.get("quality_score") or self.data.get("summary", {}).get("quality_score")
        if explicit_score is not None and int(explicit_score) > 0:
            return int(explicit_score)

        # Dynamic Auto-Calculation
        tests = self._tests()
        bugs = self._bugs()
        summary = self._summary()
        
        total_tests = len(tests)
        if not total_tests:
            total_tests = summary.get("total_tests", summary.get("total", 0))

        if total_tests > 0:
            if tests:
                passed = sum(1 for t in tests if str(t.get("status", "")).strip().lower() in ["pass", "passed", "success", "ok"])
            else:
                passed = summary.get("passed", 0)
            base_score = (passed / total_tests) * 100.0
        else:
            base_score = 100.0 if not bugs else 75.0

        # Bug deductions based on severity
        deduction = 0.0
        for b in bugs:
            sev = str(b.get("severity", "MEDIUM")).strip().upper()
            if sev in ["CRITICAL", "1", "BLOCKER"]:
                deduction += 30.0
            elif sev in ["HIGH", "2"]:
                deduction += 18.0
            elif sev in ["MEDIUM", "3"]:
                deduction += 10.0
            elif sev in ["LOW", "4", "5"]:
                deduction += 4.0

        final_score = max(0, min(100, int(round(base_score - (deduction * 0.4)))))
        return final_score

    def _embed_file(self, path_str, file_type="image"):
        """Embed file as base64 for HTML report. Supports image and video with multi-path lookup."""
        if not path_str:
            return None
        
        candidates = []
        p = Path(path_str)
        if p.is_absolute():
            candidates.append(p)
        else:
            candidates.append(self.output_dir / p)
            candidates.append(self.output_dir / "screenshots" / p.name)
            candidates.append(self.output_dir / "videos" / p.name)
            candidates.append(self.output_dir.parent / p)
            candidates.append(Path.cwd() / p)

        target_file = None
        for c in candidates:
            if c.exists() and c.is_file():
                target_file = c
                break

        if not target_file:
            return None

        ext = target_file.suffix.lower().lstrip(".")
        if file_type == "video":
            mime = {"mp4": "video/mp4", "webm": "video/webm", "ogg": "video/ogg",
                    "avi": "video/x-msvideo", "mov": "video/quicktime"}.get(ext, "video/mp4")
        else:
            mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                    "webp": "image/webp", "gif": "image/gif"}.get(ext, "image/png")
        try:
            with open(target_file, "rb") as f:
                b64 = base64.b64encode(f.read()).decode()
            return f"data:{mime};base64,{b64}"
        except Exception:
            return None

    def _find_video(self):
        """Find test recording video file in multiple possible locations and verify existence."""
        video_path = self.data.get("video_recording") or self.data.get("video")
        if video_path:
            # Check if direct path or relative path exists
            candidates = [
                Path(video_path),
                self.output_dir / video_path,
                self.output_dir / Path(video_path).name,
                self.output_dir / "videos" / Path(video_path).name,
                self.output_dir.parent / video_path,
                self.output_dir.parent / "videos" / Path(video_path).name,
            ]
            for cand in candidates:
                if cand.exists() and cand.is_file() and cand.stat().st_size > 0:
                    return str(cand)
        
        search_dirs = [
            self.output_dir,
            self.output_dir / "videos",
            self.output_dir / "test-results",
            self.output_dir.parent / "test-results",
            self.output_dir.parent / "videos",
            Path.cwd() / "test-results",
        ]
        for sdir in search_dirs:
            if sdir.exists():
                for ext in ["*.webm", "*.mp4", "*.mov"]:
                    videos = [v for v in sdir.glob(ext) if v.is_file() and v.stat().st_size > 0]
                    if videos:
                        return str(videos[0])
                    # Recursive search inside subfolders
                    videos_rec = [v for v in sdir.rglob(ext) if v.is_file() and v.stat().st_size > 0]
                    if videos_rec:
                        return str(videos_rec[0])
        return None

    # ── DOCX Generator (Mode-Aware UI Layout) ────────────────
    def generate_docx(self, filename="QA_TEST_REPORT.docx"):
        if not HAS_DOCX:
            print("Warning: python-docx not installed. Skipping DOCX generation.")
            return None

        from docx.shared import Cm, Emu, Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        doc = Document()

        # Detect testing mode: 'e2e' | 'bug_hunter' | 'api'
        mode = str(self.data.get("testing_mode", self.data.get("mode", ""))).strip().lower()
        if not mode:
            # Auto-detect from data composition
            if self.data.get("api_testing") or (not self._bugs() and any(t.get("layer", "").lower() == "api" for t in self._tests())):
                mode = "api"
            elif self._bugs() and not self._a11y() and not self._perf():
                mode = "bug_hunter"
            else:
                mode = "e2e"

        # Normalize alias
        if mode in ["api_testing", "api-test", "api_test"]:
            mode = "api"
        elif mode in ["bughunter", "bug-hunter", "bug"]:
            mode = "bug_hunter"

        # ── Page setup (Letter, margins top=90pt, bottom=50.4pt, left=57.6pt, right=57.6pt) ──
        sec = doc.sections[0]
        sec.header_distance = Pt(28)
        sec.top_margin = Pt(90)
        sec.bottom_margin = Pt(50.4)
        sec.left_margin = Pt(57.6)
        sec.right_margin = Pt(57.6)

        # ── Colors ──
        BRAND_NAVY = RGBColor(0x1B, 0x2A, 0x8A)      # #1B2A8A Primary Navy
        HEADER_BG = "1B2A8A"                         # Navy table header
        HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
        INFO_LABEL_BG = "F2F4FB"                     # Info label light tint
        LIGHT_SUBTITLE = RGBColor(0xDC, 0xE1, 0xF5)

        TEXT_RED = RGBColor(0xB3, 0x26, 0x1E)        # Severity 1 / Fail / Actual
        TEXT_ORANGE = RGBColor(0xC5, 0x6A, 0x12)     # Severity 2
        TEXT_YELLOW = RGBColor(0x8D, 0x73, 0x00)     # Severity 3
        TEXT_BLUE = RGBColor(0x18, 0x5F, 0xB4)       # Severity 4
        TEXT_GREEN = RGBColor(0x1E, 0x7A, 0x3C)      # Severity 5 / Pass / Expected

        SEV_THEME = {
            1: {"name": "Critical", "bg": "FDE7E4", "text_col": TEXT_RED},
            2: {"name": "High", "bg": "FDF0E3", "text_col": TEXT_ORANGE},
            3: {"name": "Medium", "bg": "FEF7DA", "text_col": TEXT_YELLOW},
            4: {"name": "Low", "bg": "E8F1FB", "text_col": TEXT_BLUE},
        }

        def get_sev_num(sev_val):
            s = str(sev_val).strip().lower()
            if s in ["1", "critical", "crit", "blocker"]: return 1
            if s in ["2", "high"]: return 2
            if s in ["3", "medium", "med"]: return 3
            if s in ["4", "5", "low", "trivial", "cosmetic", "info"]: return 4
            return 3

        # ── Logo Discovery ──
        logo_path = None
        candidates = [
            Path(__file__).parent.parent / "assets" / "logo.jpeg",
            Path(__file__).parent.parent / "assets" / "logo.jpg",
            Path(__file__).parent.parent / "assets" / "logo.png",
            Path(r"D:\Adila\Playwright\orca\Tamplate report BUG\WhatsApp Image 2026-09-01 at 14.23.23.jpeg"),
            self.output_dir / "logo.jpeg",
            self.output_dir / "logo.png",
        ]
        for c in candidates:
            if c.exists():
                logo_path = c
                break

        # Setup Header for all sections with Logo if found
        if logo_path:
            for s in doc.sections:
                header = s.header
                hp = header.paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                hp.paragraph_format.space_after = Pt(14)
                hrun = hp.add_run()
                hrun.add_picture(str(logo_path), height=Inches(0.72))

        # Default typography
        font_family = 'Calibri'
        style = doc.styles['Normal']
        style.font.name = font_family
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

        def set_cell_shading(cell, color_hex):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, font=font_family):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if align:
                p.alignment = align
            run = p.add_run(str(text))
            run.font.name = font
            run.font.size = Pt(size)
            run.bold = bold
            if color:
                run.font.color.rgb = color

        def add_header_row(table, texts, row_idx=0):
            for i, txt in enumerate(texts):
                cell = table.rows[row_idx].cells[i]
                set_cell_shading(cell, HEADER_BG)
                set_cell_text(cell, txt, bold=True, color=HEADER_FG, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

        def auto_width(table):
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:type="pct" w:w="5000"/>')
            tblPr.append(tblW)

        def add_h1(title_text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(title_text)
            run.font.name = font_family
            run.font.size = Pt(14 if mode == 'e2e' else 15)
            run.bold = True
            run.font.color.rgb = BRAND_NAVY
            return p

        def add_h2(subtitle_text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(subtitle_text)
            run.font.name = font_family
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = BRAND_NAVY
            return p

        summary = self._summary()
        bugs = self._bugs()
        tests = self._tests()
        perf = self._perf()
        a11y = self._a11y()
        project = self.data.get("project_name", "N/A")
        target = self.data.get("environment", self.data.get("target_url", "N/A"))
        test_date = self.data.get("test_date", datetime.now().strftime("%Y-%m-%d"))
        tester = self.data.get("tester", "QA Ultra Tester")

        total = summary.get("total_tests", summary.get("total", 0))
        if not total and tests:
            total = len(tests)
            passed = sum(1 for t in tests if str(t.get("status", "")).strip().lower() in ["pass", "passed", "success", "ok"])
            failed = sum(1 for t in tests if str(t.get("status", "")).strip().lower() in ["fail", "failed", "error", "defect"])
            flaky = sum(1 for t in tests if t.get("flaky"))
        else:
            passed = summary.get("passed", 0)
            failed = summary.get("failed", 0)
            flaky = summary.get("flaky", 0)
        
        quality = summary.get("quality_score", self.data.get("quality_score", 0))
        success_rate = f"{passed*100//total}%" if total else "0%"

        sev_counts = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
        for b in bugs:
            n = get_sev_num(b.get("severity", 3))
            sev_counts[n] = sev_counts.get(n, 0) + 1

        if sev_counts[1] > 0:
            readiness = "BELUM SIAP – Ada Isu Severity 1 (Critical)"
            readiness_col = TEXT_RED
        elif sev_counts[2] > 0:
            readiness = "BERSYARAT – Ada Isu Severity 2 (High)"
            readiness_col = TEXT_ORANGE
        elif failed > 0:
            readiness = "LULUS DENGAN CATATAN"
            readiness_col = TEXT_YELLOW
        else:
            readiness = "SIAP RILIS (PRODUCTION READY)"
            readiness_col = TEXT_GREEN

        # ════════════════════════════════════════════════════════
        # ROUTE LAYOUT SESUAI MODE
        # ════════════════════════════════════════════════════════

        # ── MODE 1: TESTING E2E (Formal Comprehensive Layout) ───
        if mode == "e2e":
            # HEADER BANNER (Single cell with Navy BG - Consistent with Bug Hunter & API)
            banner_tbl = doc.add_table(rows=1, cols=1)
            banner_tbl.style = 'Table Grid'
            auto_width(banner_tbl)
            b_cell = banner_tbl.rows[0].cells[0]
            set_cell_shading(b_cell, HEADER_BG)
            
            bp0 = b_cell.paragraphs[0]
            bp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp0.paragraph_format.space_before = Pt(8)
            bp0.paragraph_format.space_after = Pt(2)
            r0 = bp0.add_run("LAPORAN HASIL PENGUJIAN QA E2E")
            r0.font.name = 'Calibri'
            r0.font.size = Pt(22)
            r0.bold = True
            r0.font.color.rgb = HEADER_FG

            bp1 = b_cell.add_paragraph()
            bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp1.paragraph_format.space_before = Pt(0)
            bp1.paragraph_format.space_after = Pt(8)
            r1 = bp1.add_run(f"END-TO-END AUTOMATION REPORT – {project.upper()}")
            r1.font.name = 'Calibri'
            r1.font.size = Pt(13)
            r1.bold = True
            r1.font.color.rgb = LIGHT_SUBTITLE

            doc.add_paragraph('')

            # Info Table
            meta_tbl = doc.add_table(rows=6, cols=2)
            meta_tbl.style = 'Table Grid'
            auto_width(meta_tbl)
            meta_data = [
                ("Target / Environment", target),
                ("Fitur / Modul Uji", project),
                ("Tanggal Uji", test_date),
                ("Metodologi / Tester", tester),
                ("Quality Score", f"{quality}/100" if quality else "-"),
                ("Status Kesiapan", readiness),
            ]
            for i, (label, val) in enumerate(meta_data):
                set_cell_shading(meta_tbl.rows[i].cells[0], INFO_LABEL_BG)
                set_cell_text(meta_tbl.rows[i].cells[0], label, bold=True, size=10)
                if label == "Status Kesiapan":
                    set_cell_text(meta_tbl.rows[i].cells[1], val, bold=True, color=readiness_col, size=10)
                else:
                    set_cell_text(meta_tbl.rows[i].cells[1], val, size=10)

            # Section 1: Executive Summary
            add_h1("1. Ringkasan Eksekutif (Executive Summary)")
            narrative = (
                f"Pengujian end-to-end (E2E) telah selesai dieksekusi pada {project}. "
                f"Dari total {total} skenario, {passed} skenario PASSED dan {failed} skenario FAILED "
                f"dengan tingkat keberhasilan {success_rate}."
            )
            if bugs:
                narrative += f" Ditemukan total {len(bugs)} defect/bug selama pengujian."
            p = doc.add_paragraph(narrative)
            p.paragraph_format.space_after = Pt(6)

            sum_tbl = doc.add_table(rows=2, cols=4)
            sum_tbl.style = 'Table Grid'
            auto_width(sum_tbl)
            add_header_row(sum_tbl, ["Total Skenario", "Passed", "Failed", "Success Rate"])
            set_cell_text(sum_tbl.rows[1].cells[0], str(total), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(sum_tbl.rows[1].cells[1], str(passed), color=TEXT_GREEN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(sum_tbl.rows[1].cells[2], str(failed), color=TEXT_RED if failed > 0 else TEXT_GREEN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(sum_tbl.rows[1].cells[3], success_rate, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

            # Statistik Temuan Bug (Severity 1–4)
            sev_counts = {1: 0, 2: 0, 3: 0, 4: 0}
            for bug in bugs:
                s_num = get_sev_num(bug.get("severity", 3))
                sev_counts[s_num] = sev_counts.get(s_num, 0) + 1

            p_sev_title = doc.add_paragraph()
            p_sev_title.paragraph_format.space_before = Pt(8)
            p_sev_title.paragraph_format.space_after = Pt(4)
            r_st = p_sev_title.add_run("Statistik Temuan Defect Berdasarkan Tingkat Keparahan (Severity 1–4):")
            r_st.bold = True
            r_st.font.name = font_family
            r_st.font.size = Pt(10.5)

            sev_tbl = doc.add_table(rows=5, cols=3)
            sev_tbl.style = 'Table Grid'
            auto_width(sev_tbl)
            add_header_row(sev_tbl, ["Severity", "Keterangan", "Jumlah"])
            for n in range(1, 5):
                theme = SEV_THEME[n]
                set_cell_shading(sev_tbl.rows[n].cells[0], theme["bg"])
                set_cell_text(sev_tbl.rows[n].cells[0], str(n), bold=True, color=theme["text_col"], align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_shading(sev_tbl.rows[n].cells[1], theme["bg"])
                set_cell_text(sev_tbl.rows[n].cells[1], theme["name"], color=theme["text_col"])
                set_cell_text(sev_tbl.rows[n].cells[2], str(sev_counts[n]), align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph('')

            # Section 2: Hasil Pengujian Fitur & Skenario (Happy Path & Negative)
            add_h1("2. Hasil Pengujian Skenario (Happy Path & Negative)")
            
            happy_tests = [t for t in tests if t.get("category") == "happy_path" or (not t.get("category") and str(t.get("status","")).strip().lower() in ["pass", "passed", "success", "ok"] and "invalid" not in t.get("name","").lower() and "fail" not in t.get("name","").lower())]
            negative_tests = [t for t in tests if t not in happy_tests]

            def _render_test_scenario_tables(test_list, category_title, start_num):
                if not test_list:
                    return start_num
                
                h2 = doc.add_paragraph()
                h2.paragraph_format.space_before = Pt(8)
                h2.paragraph_format.space_after = Pt(4)
                r_h2 = h2.add_run(category_title)
                r_h2.font.name = 'Calibri'
                r_h2.font.size = Pt(12)
                r_h2.bold = True
                r_h2.font.color.rgb = BRAND_NAVY

                for idx, test in enumerate(test_list, start_num):
                    status = str(test.get("status", "")).strip().lower()
                    is_pass = status in ["pass", "passed", "success", "ok"]
                    status_text = "PASSED" if is_pass else "FAILED"
                    status_color = TEXT_GREEN if is_pass else TEXT_RED
                    duration_text = format_duration(test.get("duration_ms", test.get("duration", 0)))
                    
                    steps = test.get("steps", [])
                    if isinstance(steps, list) and steps:
                        steps_str = "\n".join([f"{strip_step_number(s)}" for s in steps])
                    elif isinstance(steps, str) and steps:
                        steps_str = steps
                    else:
                        steps_str = "-"
                        
                    expected_str = test.get("expected", "-")
                    actual_str = test.get("actual", test.get("details", test.get("error", "-")))

                    # Scenario detail card table
                    card_tbl = doc.add_table(rows=5, cols=2)
                    card_tbl.style = 'Table Grid'
                    auto_width(card_tbl)
                    
                    # Row 0: Header with Test Name & Status
                    set_cell_shading(card_tbl.rows[0].cells[0], INFO_LABEL_BG)
                    set_cell_shading(card_tbl.rows[0].cells[1], INFO_LABEL_BG)
                    set_cell_text(card_tbl.rows[0].cells[0], f"Skenario #{idx}: {test.get('name', '')}", bold=True, color=BRAND_NAVY, size=10)
                    set_cell_text(card_tbl.rows[0].cells[1], f"[{status_text}] ({duration_text})", bold=True, color=status_color, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

                    # Row 1: Steps to Reproduce
                    set_cell_shading(card_tbl.rows[1].cells[0], INFO_LABEL_BG)
                    set_cell_text(card_tbl.rows[1].cells[0], "Steps / Langkah", bold=True, size=9)
                    set_cell_text(card_tbl.rows[1].cells[1], steps_str, size=9)

                    # Row 2: Expected Result
                    set_cell_shading(card_tbl.rows[2].cells[0], INFO_LABEL_BG)
                    set_cell_text(card_tbl.rows[2].cells[0], "Expected Result", bold=True, size=9)
                    set_cell_text(card_tbl.rows[2].cells[1], expected_str, color=TEXT_GREEN, size=9)

                    # Row 3: Actual Result
                    set_cell_shading(card_tbl.rows[3].cells[0], INFO_LABEL_BG)
                    set_cell_text(card_tbl.rows[3].cells[0], "Actual Result", bold=True, size=9)
                    set_cell_text(card_tbl.rows[3].cells[1], actual_str, color=TEXT_RED if not is_pass else None, size=9)

                    # Row 4: Layer & Environment
                    set_cell_shading(card_tbl.rows[4].cells[0], INFO_LABEL_BG)
                    set_cell_text(card_tbl.rows[4].cells[0], "Layer / Browser", bold=True, size=9)
                    set_cell_text(card_tbl.rows[4].cells[1], f"{test.get('layer', 'UI').upper()} | {test.get('browser', 'Chromium')}", size=9)

                    doc.add_paragraph('').paragraph_format.space_after = Pt(4)

                return start_num + len(test_list)

            curr_idx = 1
            if happy_tests:
                curr_idx = _render_test_scenario_tables(happy_tests, f"2.1 Skenario Happy Path ({len(happy_tests)} Tests)", curr_idx)
            if negative_tests:
                curr_idx = _render_test_scenario_tables(negative_tests, f"2.2 Skenario Negative / Boundary ({len(negative_tests)} Tests)", curr_idx)
            if not tests:
                doc.add_paragraph("Tidak ada detail skenario test yang tercatat.")

            # Section 3: Performance Core Web Vitals
            if perf:
                add_h1("3. Metrik Performa (Core Web Vitals & Kecepatan Akses)")
                p_perf = doc.add_paragraph("Evaluasi performa halaman sesuai standar industri Google Core Web Vitals untuk memastikan kelancaran loading dan responsivitas interaksi pengguna:")
                p_perf.paragraph_format.space_after = Pt(4)

                perf_items = []
                for key, val in perf.items():
                    metric_clean = key.replace("_ms", "").replace("_", " ").strip().upper()
                    key_map = {
                        "PAGE LOAD MS": "PAGE LOAD",
                        "PAGE LOAD": "PAGE LOAD",
                        "TTFB MS": "TTFB",
                        "FCP MS": "FCP",
                        "LCP MS": "LCP",
                        "CLS": "CLS",
                        "FID": "FID",
                        "INP": "INP",
                        "TBT": "TBT"
                    }
                    display_key = key_map.get(metric_clean, metric_clean)
                    info = METRIC_INFO.get(display_key, METRIC_INFO.get(metric_clean, {}))
                    full_name = info.get("full", key)
                    desc_text = info.get("desc", "")
                    
                    if isinstance(val, dict):
                        val_text = f'{val.get("value","")}{val.get("unit","")}'
                        rating = val.get("rating", "")
                    elif "ms" in key.lower():
                        val_text = format_duration(val) if isinstance(val, (int, float)) else str(val)
                        rating = ""
                    else:
                        val_text = str(val)
                        rating = ""
                    perf_items.append((display_key, full_name, desc_text, val_text, rating))

                perf_tbl = doc.add_table(rows=len(perf_items) + 1, cols=4)
                perf_tbl.style = 'Table Grid'
                auto_width(perf_tbl)
                add_header_row(perf_tbl, ["Metrik & Penjelasan", "Nama Standar", "Nilai Uji", "Rating Kelayakan"])
                for idx, (metric, full, desc, val_t, rating) in enumerate(perf_items, 1):
                    if idx % 2 == 0:
                        for cell in perf_tbl.rows[idx].cells: set_cell_shading(cell, INFO_LABEL_BG)
                    
                    # Cell 0: Metric + Deskripsi
                    c0 = perf_tbl.rows[idx].cells[0]
                    c0.text = ""
                    p0 = c0.paragraphs[0]
                    p0.paragraph_format.space_before = Pt(2)
                    p0.paragraph_format.space_after = Pt(1)
                    r_m = p0.add_run(metric)
                    r_m.bold = True
                    r_m.font.name = font_family
                    r_m.font.size = Pt(10)
                    if desc:
                        p0_sub = c0.add_paragraph()
                        p0_sub.paragraph_format.space_before = Pt(0)
                        p0_sub.paragraph_format.space_after = Pt(2)
                        r_d = p0_sub.add_run(desc)
                        r_d.italic = True
                        r_d.font.name = font_family
                        r_d.font.size = Pt(8.5)
                        r_d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    set_cell_text(perf_tbl.rows[idx].cells[1], full, size=9.5)
                    set_cell_text(perf_tbl.rows[idx].cells[2], val_t, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
                    
                    r_col = TEXT_GREEN if rating.lower() in ["good", "optimal", "pass"] else (TEXT_ORANGE if rating.lower() in ["needs_improvement", "warning"] else (TEXT_RED if rating.lower() in ["poor", "bad", "fail"] else None))
                    set_cell_text(perf_tbl.rows[idx].cells[3], rating.replace("_", " ").title() if rating else "-", bold=True if r_col else False, color=r_col, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
                doc.add_paragraph('')

            # Section 4: Accessibility
            if a11y:
                add_h1("4. Aksesibilitas & Kenyamanan Pengguna (WCAG 2.1 AA)")
                p_a11y = doc.add_paragraph("Pemeriksaan terhadap standar kenyamanan antarmuka web (kemudahan membaca teks, kontras warna, label form, dan aksesibilitas):")
                p_a11y.paragraph_format.space_after = Pt(4)
                
                a11y_tbl = doc.add_table(rows=len(a11y) + 1, cols=4)
                a11y_tbl.style = 'Table Grid'
                auto_width(a11y_tbl)
                add_header_row(a11y_tbl, ["Dampak", "Jenis Temuan & Penjelasan", "Elemen UI", "Standar Acuan"])
                for idx, issue in enumerate(a11y, 1):
                    impact = issue.get("impact", issue.get("severity", "moderate")).lower()
                    impact_color = TEXT_RED if impact in ["critical", "serious"] else (TEXT_ORANGE if impact == "moderate" else None)
                    rule = issue.get("type", issue.get("rule", "Aksesibilitas UI"))
                    desc = issue.get("description", "")
                    wcag_raw = issue.get("wcag", "")

                    wcag_map = {
                        "1.1.1": "1.1.1 Non-text Content (Teks Alternatif Gambar)",
                        "1.3.1": "1.3.1 Info & Relationships (Struktur Form & Label)",
                        "1.4.3": "1.4.3 Contrast Minimum (Kontras Warna Teks)",
                        "2.1.1": "2.1.1 Keyboard Accessible (Akses Tombol/Keyboard)",
                        "2.4.7": "2.4.7 Focus Visible (Indikator Fokus Terlihat)",
                        "3.3.2": "3.3.2 Labels or Instructions (Petunjuk Pengisian)"
                    }
                    wcag_friendly = wcag_raw
                    for k, v in wcag_map.items():
                        if k in wcag_raw:
                            wcag_friendly = f"WCAG 2.1 AA — {v}"
                            break

                    set_cell_text(a11y_tbl.rows[idx].cells[0], impact.upper(), color=impact_color, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                    
                    # Cell 1: Rule + Desc
                    c_desc = a11y_tbl.rows[idx].cells[1]
                    c_desc.text = ""
                    p1 = c_desc.paragraphs[0]
                    r_r = p1.add_run(f"{rule}\n")
                    r_r.bold = True
                    r_r.font.name = font_family
                    r_r.font.size = Pt(9.5)
                    r_d = p1.add_run(desc)
                    r_d.font.name = font_family
                    r_d.font.size = Pt(9)
                    r_d.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

                    set_cell_text(a11y_tbl.rows[idx].cells[2], issue.get("element", "-"), size=9)
                    set_cell_text(a11y_tbl.rows[idx].cells[3], wcag_friendly, bold=True, color=BRAND_NAVY, size=9)

            # Section 5: Daftar Defect jika ada
            if bugs:
                add_h1("5. Ringkasan Defect / Bug")
                bug_tbl = doc.add_table(rows=len(bugs) + 1, cols=5)
                bug_tbl.style = 'Table Grid'
                auto_width(bug_tbl)
                add_header_row(bug_tbl, ["ID", "Judul", "Severity", "Priority", "Status"])
                for idx, bug in enumerate(bugs, 1):
                    s_num = get_sev_num(bug.get("severity", 3))
                    theme = SEV_THEME[s_num]
                    set_cell_text(bug_tbl.rows[idx].cells[0], bug.get("id", f"BUG-{idx:03d}"), bold=True)
                    set_cell_text(bug_tbl.rows[idx].cells[1], bug.get("title", ""))
                    set_cell_shading(bug_tbl.rows[idx].cells[2], theme["bg"])
                    set_cell_text(bug_tbl.rows[idx].cells[2], theme["name"], bold=True, color=theme["text_col"])
                    set_cell_text(bug_tbl.rows[idx].cells[3], bug.get("priority", "-"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_text(bug_tbl.rows[idx].cells[4], bug.get("status", "Open"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_text(bug_tbl.rows[idx].cells[0], bug.get("id", f"BUG-{idx:03d}"), bold=True)
                    set_cell_text(bug_tbl.rows[idx].cells[1], bug.get("title", ""))
                    set_cell_shading(bug_tbl.rows[idx].cells[2], theme["bg"])
                    set_cell_text(bug_tbl.rows[idx].cells[2], theme["name"], bold=True, color=theme["text_col"])
                    set_cell_text(bug_tbl.rows[idx].cells[3], bug.get("priority", "-"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_text(bug_tbl.rows[idx].cells[4], bug.get("status", "Open"), align=WD_ALIGN_PARAGRAPH.CENTER)

        # ── MODE 2: BUG HUNTER (Branded Compact & Evidence Centric) ───
        elif mode == "bug_hunter":
            # HEADER BANNER (Single cell with Navy BG)
            banner_tbl = doc.add_table(rows=1, cols=1)
            banner_tbl.style = 'Table Grid'
            auto_width(banner_tbl)
            b_cell = banner_tbl.rows[0].cells[0]
            set_cell_shading(b_cell, HEADER_BG)
            
            bp0 = b_cell.paragraphs[0]
            bp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp0.paragraph_format.space_before = Pt(8)
            bp0.paragraph_format.space_after = Pt(2)
            r0 = bp0.add_run("LAPORAN TEMUAN BUG & DEFECT AUDIT")
            r0.font.name = 'Calibri'
            r0.font.size = Pt(22)
            r0.bold = True
            r0.font.color.rgb = HEADER_FG

            bp1 = b_cell.add_paragraph()
            bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp1.paragraph_format.space_before = Pt(0)
            bp1.paragraph_format.space_after = Pt(8)
            r1 = bp1.add_run(f"BUG HUNTER REPORT – {project.upper()}")
            r1.font.name = 'Calibri'
            r1.font.size = Pt(13)
            r1.bold = True
            r1.font.color.rgb = LIGHT_SUBTITLE

            doc.add_paragraph('')

            # Info Table
            meta_tbl = doc.add_table(rows=5, cols=2)
            meta_tbl.style = 'Table Grid'
            auto_width(meta_tbl)
            meta_data = [
                ("Target / Environment", target),
                ("Fitur / Menu Uji", project),
                ("Tanggal Audit", test_date),
                ("Metodologi / Hunter", tester),
                ("Status Kesiapan", readiness),
            ]
            for i, (label, val) in enumerate(meta_data):
                set_cell_shading(meta_tbl.rows[i].cells[0], INFO_LABEL_BG)
                set_cell_text(meta_tbl.rows[i].cells[0], label, bold=True, size=10)
                if label == "Status Kesiapan":
                    set_cell_text(meta_tbl.rows[i].cells[1], val, bold=True, color=readiness_col, size=10)
                else:
                    set_cell_text(meta_tbl.rows[i].cells[1], val, size=10)

            # 1. Ringkasan Eksekutif
            add_h1("1. Ringkasan Eksekutif")
            sev_desc_parts = []
            for n in range(1, 5):
                if sev_counts[n] > 0:
                    sev_desc_parts.append(f"{sev_counts[n]} {SEV_THEME[n]['name']}")
            sev_narrative = f" Ditemukan {len(bugs)} temuan ({', '.join(sev_desc_parts)})." if bugs else " Tidak ditemukan defect/bug."
            p = doc.add_paragraph(f"Audit Bug Hunting telah selesai dieksekusi pada fitur {project}.{sev_narrative}")
            p.paragraph_format.space_after = Pt(6)

            # 2. Statistik Temuan (Severity 1–4)
            add_h1("2. Statistik Temuan (Severity 1–4)")
            sev_tbl = doc.add_table(rows=5, cols=3)
            sev_tbl.style = 'Table Grid'
            auto_width(sev_tbl)
            add_header_row(sev_tbl, ["Severity", "Keterangan", "Jumlah"])
            for n in range(1, 5):
                theme = SEV_THEME[n]
                set_cell_shading(sev_tbl.rows[n].cells[0], theme["bg"])
                set_cell_text(sev_tbl.rows[n].cells[0], str(n), bold=True, color=theme["text_col"], align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_shading(sev_tbl.rows[n].cells[1], theme["bg"])
                set_cell_text(sev_tbl.rows[n].cells[1], theme["name"], color=theme["text_col"])
                set_cell_text(sev_tbl.rows[n].cells[2], str(sev_counts[n]), align=WD_ALIGN_PARAGRAPH.CENTER)

            # 3. Rincian Temuan Bug
            add_h1("3. Rincian Temuan Bug")
            if bugs:
                bug_overview = doc.add_table(rows=len(bugs) + 1, cols=5)
                bug_overview.style = 'Table Grid'
                auto_width(bug_overview)
                add_header_row(bug_overview, ["ID", "Judul", "Severity", "Priority", "Status"])
                for idx, bug in enumerate(bugs, 1):
                    s_num = get_sev_num(bug.get("severity", 3))
                    theme = SEV_THEME[s_num]
                    set_cell_text(bug_overview.rows[idx].cells[0], bug.get("id", f"BUG-{idx:03d}"), bold=True)
                    set_cell_text(bug_overview.rows[idx].cells[1], bug.get("title", ""))
                    set_cell_shading(bug_overview.rows[idx].cells[2], theme["bg"])
                    set_cell_text(bug_overview.rows[idx].cells[2], theme["name"], bold=True, color=theme["text_col"])
                    set_cell_text(bug_overview.rows[idx].cells[3], bug.get("priority", "-"), align=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_text(bug_overview.rows[idx].cells[4], bug.get("status", "Open"), align=WD_ALIGN_PARAGRAPH.CENTER)

                for bug in bugs:
                    s_num = get_sev_num(bug.get("severity", 3))
                    theme = SEV_THEME[s_num]
                    b_id = bug.get("id", "BUG")
                    b_title = bug.get("title", "")
                    add_h2(f"{b_id} – {b_title}")

                    # Structured Table for Bug Details
                    detail_tbl = doc.add_table(rows=8, cols=2)
                    detail_tbl.style = 'Table Grid'
                    auto_width(detail_tbl)

                    # Row 0: Severity
                    set_cell_shading(detail_tbl.rows[0].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[0].cells[0], "Severity", bold=True, size=10)
                    set_cell_shading(detail_tbl.rows[0].cells[1], theme["bg"])
                    set_cell_text(detail_tbl.rows[0].cells[1], f"Severity {s_num} – {theme['name']}", bold=True, color=theme["text_col"], size=10)

                    # Row 1: Priority
                    set_cell_shading(detail_tbl.rows[1].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[1].cells[0], "Priority", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[1].cells[1], bug.get("priority", "-"), bold=True, size=10)

                    # Row 2: Tipe Bug
                    set_cell_shading(detail_tbl.rows[2].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[2].cells[0], "Tipe Bug", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[2].cells[1], bug.get("type", "-"), size=10)

                    # Row 3: Lokasi
                    set_cell_shading(detail_tbl.rows[3].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[3].cells[0], "Lokasi / Endpoint", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[3].cells[1], bug.get("location", target), size=10)

                    # Row 4: Langkah Reproduksi
                    set_cell_shading(detail_tbl.rows[4].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[4].cells[0], "Langkah Reproduksi", bold=True, size=10)
                    steps = bug.get("steps", [])
                    c_step = detail_tbl.rows[4].cells[1]
                    c_step.text = ""
                    if steps:
                        for idx_s, step in enumerate(steps, 1):
                            clean = strip_step_number(step)
                            p_st = c_step.paragraphs[0] if idx_s == 1 else c_step.add_paragraph()
                            p_st.paragraph_format.space_before = Pt(1)
                            p_st.paragraph_format.space_after = Pt(1)
                            r_n = p_st.add_run(f"{idx_s}. ")
                            r_n.bold = True
                            r_n.font.name = font_family
                            r_n.font.size = Pt(10)
                            r_txt = p_st.add_run(clean)
                            r_txt.font.name = font_family
                            r_txt.font.size = Pt(10)
                    else:
                        set_cell_text(c_step, "-", size=10)

                    # Row 5: Hasil Aktual
                    set_cell_shading(detail_tbl.rows[5].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[5].cells[0], "Hasil Aktual", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[5].cells[1], bug.get("actual", "-"), color=TEXT_RED, size=10)

                    # Row 6: Hasil Diharapkan
                    set_cell_shading(detail_tbl.rows[6].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[6].cells[0], "Hasil Diharapkan", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[6].cells[1], bug.get("expected", "-"), color=TEXT_GREEN, size=10)

                    # Row 7: Rekomendasi
                    set_cell_shading(detail_tbl.rows[7].cells[0], INFO_LABEL_BG)
                    set_cell_text(detail_tbl.rows[7].cells[0], "Rekomendasi", bold=True, size=10)
                    set_cell_text(detail_tbl.rows[7].cells[1], bug.get("recommendation", "-"), bold=True, color=BRAND_NAVY, size=10)

                    evidence = bug.get("evidence", "")
                    screenshot = evidence if isinstance(evidence, str) else ""
                    if not screenshot and isinstance(evidence, dict):
                        screenshot = evidence.get("screenshot", "")
                    if screenshot:
                        spath = Path(screenshot)
                        if not spath.is_absolute():
                            spath = self.output_dir / spath
                        if spath.exists():
                            try:
                                doc.add_paragraph('')
                                doc.add_picture(str(spath), width=Inches(5.5))
                                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cap = doc.add_paragraph()
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                r_cap = cap.add_run(f'Gambar: Bukti {b_id} – {b_title[:60]}')
                                r_cap.font.size = Pt(9)
                                r_cap.italic = True
                                r_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                            except Exception as e:
                                print(f"Warning: Could not embed screenshot: {e}")
                    doc.add_paragraph('')

            # 4. Rekomendasi Tindak Lanjut
            add_h1("4. Rekomendasi Tindak Lanjut")
            if bugs:
                for idx_b, bug in enumerate(bugs, 1):
                    s_num = get_sev_num(bug.get("severity", 3))
                    theme = SEV_THEME[s_num]
                    rec = bug.get("recommendation", "Perlu investigasi dan perbaikan lebih lanjut.")
                    p_rec = doc.add_paragraph()
                    r_prefix = p_rec.add_run(f"{idx_b}. [Severity {s_num}] {bug.get('id','BUG')} – ")
                    r_prefix.bold = True
                    r_prefix.font.color.rgb = theme["text_col"]
                    p_rec.add_run(rec)

        # ── MODE 3: API TESTING (Technical Endpoint & Contract Layout) ───
        elif mode == "api":
            # HEADER BANNER (Technical Style)
            banner_tbl = doc.add_table(rows=1, cols=1)
            banner_tbl.style = 'Table Grid'
            auto_width(banner_tbl)
            b_cell = banner_tbl.rows[0].cells[0]
            set_cell_shading(b_cell, HEADER_BG)
            
            bp0 = b_cell.paragraphs[0]
            bp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp0.paragraph_format.space_before = Pt(8)
            bp0.paragraph_format.space_after = Pt(2)
            r0 = bp0.add_run("LAPORAN PENGUJIAN API & INTEGRASI")
            r0.font.name = 'Calibri'
            r0.font.size = Pt(20)
            r0.bold = True
            r0.font.color.rgb = HEADER_FG

            bp1 = b_cell.add_paragraph()
            bp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp1.paragraph_format.space_before = Pt(0)
            bp1.paragraph_format.space_after = Pt(8)
            r1 = bp1.add_run(f"API VALIDATION & CONTRACT REPORT – {project.upper()}")
            r1.font.name = 'Calibri'
            r1.font.size = Pt(12)
            r1.bold = True
            r1.font.color.rgb = LIGHT_SUBTITLE

            doc.add_paragraph('')

            # Info Table
            meta_tbl = doc.add_table(rows=5, cols=2)
            meta_tbl.style = 'Table Grid'
            auto_width(meta_tbl)
            meta_data = [
                ("Base URL / Endpoint", target),
                ("Service / Module", project),
                ("Tanggal Uji", test_date),
                ("Tester / Tool", tester),
                ("Status API", "LULUS SEMUA" if failed == 0 else f"ADA {failed} ENDPOINT ERROR"),
            ]
            for i, (label, val) in enumerate(meta_data):
                set_cell_shading(meta_tbl.rows[i].cells[0], INFO_LABEL_BG)
                set_cell_text(meta_tbl.rows[i].cells[0], label, bold=True, size=10)
                if label == "Status API":
                    set_cell_text(meta_tbl.rows[i].cells[1], val, bold=True, color=TEXT_GREEN if failed == 0 else TEXT_RED, size=10)
                else:
                    set_cell_text(meta_tbl.rows[i].cells[1], val, size=10)

            # 1. API Execution Matrix
            add_h1("1. Matriks Pengujian Endpoint API")
            if tests:
                api_tbl = doc.add_table(rows=len(tests) + 1, cols=5)
                api_tbl.style = 'Table Grid'
                auto_width(api_tbl)
                add_header_row(api_tbl, ["No", "Endpoint & Method", "Expected Result", "Status", "Latency"])
                for i, t in enumerate(tests):
                    row = api_tbl.rows[i + 1]
                    if i % 2 == 1:
                        for cell in row.cells: set_cell_shading(cell, INFO_LABEL_BG)
                    st = str(t.get("status", "")).strip().lower()
                    is_p = st in ["pass", "passed", "success", "ok"]
                    meth = t.get("method", "POST" if "create" in t.get("name","").lower() or "submit" in t.get("name","").lower() else "GET")
                    ep = t.get("endpoint", t.get("location", t.get("name", "-")))
                    dur = t.get("duration_ms", t.get("duration", 0))

                    set_cell_text(row.cells[0], str(i + 1), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
                    set_cell_text(row.cells[1], f"[{meth}] {ep}", bold=True, size=9)
                    set_cell_text(row.cells[2], t.get("expected", "HTTP 200 / Valid schema"), size=9)
                    set_cell_text(row.cells[3], "PASS" if is_p else "FAIL", bold=True, color=TEXT_GREEN if is_p else TEXT_RED, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
                    set_cell_text(row.cells[4], format_duration(dur), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

            # 2. Performance & Latency (if available)
            if perf:
                add_h1("2. Metrik Performa & Latensi API")
                p_perf = doc.add_paragraph("Pengukuran kecepatan respon server, latensi pengolahan data, dan stabilitas throughput service API:")
                p_perf.paragraph_format.space_after = Pt(4)

                perf_items = []
                for key, val in perf.items():
                    metric_clean = key.upper().replace("_", " ").strip()
                    key_map = {
                        "AVG RESPONSE TIME": "AVG RESPONSE TIME",
                        "AVERAGE RESPONSE TIME": "AVG RESPONSE TIME",
                        "P95 LATENCY": "P95 LATENCY",
                        "P95": "P95 LATENCY",
                        "THROUGHPUT": "THROUGHPUT",
                        "RPS": "THROUGHPUT",
                        "TTFB MS": "TTFB",
                        "TTFB": "TTFB"
                    }
                    display_key = key_map.get(metric_clean, metric_clean)
                    info = METRIC_INFO.get(display_key, METRIC_INFO.get(metric_clean, {}))
                    full_name = info.get("full", key)
                    desc_text = info.get("desc", "")

                    if isinstance(val, dict):
                        val_text = f'{val.get("value","")}{val.get("unit","")}'
                        rating = val.get("rating", "")
                    elif "ms" in key.lower() or "latency" in key.lower() or "time" in key.lower():
                        val_text = format_duration(val) if isinstance(val, (int, float)) else str(val)
                        rating = ""
                    else:
                        val_text = str(val)
                        rating = ""
                    perf_items.append((display_key, full_name, desc_text, val_text, rating))

                perf_tbl = doc.add_table(rows=len(perf_items) + 1, cols=4)
                perf_tbl.style = 'Table Grid'
                auto_width(perf_tbl)
                add_header_row(perf_tbl, ["Metrik & Penjelasan", "Nama Standar", "Nilai Uji", "Rating Kelayakan"])
                for idx, (metric, full, desc, val_t, rating) in enumerate(perf_items, 1):
                    if idx % 2 == 0:
                        for cell in perf_tbl.rows[idx].cells: set_cell_shading(cell, INFO_LABEL_BG)

                    c0 = perf_tbl.rows[idx].cells[0]
                    c0.text = ""
                    p0 = c0.paragraphs[0]
                    p0.paragraph_format.space_before = Pt(2)
                    p0.paragraph_format.space_after = Pt(1)
                    r_m = p0.add_run(metric)
                    r_m.bold = True
                    r_m.font.name = font_family
                    r_m.font.size = Pt(10)
                    if desc:
                        p0_sub = c0.add_paragraph()
                        p0_sub.paragraph_format.space_before = Pt(0)
                        p0_sub.paragraph_format.space_after = Pt(2)
                        r_d = p0_sub.add_run(desc)
                        r_d.italic = True
                        r_d.font.name = font_family
                        r_d.font.size = Pt(8.5)
                        r_d.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    set_cell_text(perf_tbl.rows[idx].cells[1], full, size=9.5)
                    set_cell_text(perf_tbl.rows[idx].cells[2], val_t, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

                    r_col = TEXT_GREEN if rating.lower() in ["good", "optimal", "pass"] else (TEXT_ORANGE if rating.lower() in ["needs_improvement", "warning"] else (TEXT_RED if rating.lower() in ["poor", "bad", "fail"] else None))
                    set_cell_text(perf_tbl.rows[idx].cells[3], rating.replace("_", " ").title() if rating else "-", bold=True if r_col else False, color=r_col, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
                doc.add_paragraph('')

            # 3. Bug / Defect Temuan di API
            if bugs:
                add_h1("3. Temuan Isu & Bug API")
                for bug in bugs:
                    b_id = bug.get("id", "BUG")
                    b_title = bug.get("title", "")
                    s_num = get_sev_num(bug.get("severity", 3))
                    theme = SEV_THEME[s_num]

                    add_h2(f"[{b_id}] {b_title}")
                    b_tbl = doc.add_table(rows=5, cols=2)
                    b_tbl.style = 'Table Grid'
                    auto_width(b_tbl)
                    b_rows = [
                        ("Severity", f"Severity {s_num} - {theme['name']}"),
                        ("Endpoint / Method", bug.get("location", "-")),
                        ("Hasil Aktual", bug.get("actual", "-")),
                        ("Hasil Diharapkan", bug.get("expected", "-")),
                        ("Rekomendasi Fix", bug.get("recommendation", "-")),
                    ]
                    for r_idx, (lbl, val) in enumerate(b_rows):
                        set_cell_shading(b_tbl.rows[r_idx].cells[0], INFO_LABEL_BG)
                        set_cell_text(b_tbl.rows[r_idx].cells[0], lbl, bold=True, size=10)
                        if lbl == "Severity":
                            set_cell_shading(b_tbl.rows[r_idx].cells[1], theme["bg"])
                            set_cell_text(b_tbl.rows[r_idx].cells[1], val, bold=True, color=theme["text_col"], size=10)
                        elif lbl == "Hasil Aktual":
                            set_cell_text(b_tbl.rows[r_idx].cells[1], val, color=TEXT_RED, size=10)
                        elif lbl == "Hasil Diharapkan":
                            set_cell_text(b_tbl.rows[r_idx].cells[1], val, color=TEXT_GREEN, size=10)
                        elif lbl == "Rekomendasi Fix":
                            set_cell_text(b_tbl.rows[r_idx].cells[1], val, bold=True, color=BRAND_NAVY, size=10)
                        else:
                            set_cell_text(b_tbl.rows[r_idx].cells[1], val, size=10)
                    doc.add_paragraph('')

        # ── COMMON FOOTER: BUKTI VISUAL (Gallery) ─────────────────
        ss_dir = self.output_dir / "screenshots"
        screenshots = []
        if ss_dir.exists():
            for ext in ["*.png", "*.jpg", "*.jpeg"]:
                screenshots.extend(sorted(ss_dir.glob(ext)))

        if screenshots and mode != "bug_hunter":  # Bug hunter already embeds inside bug items
            # Map screenshots to tests / bugs for richer descriptions
            ss_map = {}
            for t in tests:
                t_ss = t.get("screenshot") or ""
                if isinstance(t.get("evidence"), dict):
                    t_ss = t_ss or t["evidence"].get("screenshot", "")
                if t_ss:
                    pname = Path(t_ss).name.lower()
                    ss_map[pname] = {
                        "type": "Test Skenario",
                        "title": t.get("name", ""),
                        "status": t.get("status", "PASS"),
                        "details": t.get("actual") or t.get("details") or t.get("expected") or ""
                    }
            for b in bugs:
                b_ev = b.get("evidence") or {}
                b_ss = b_ev.get("screenshot") if isinstance(b_ev, dict) else ""
                if not b_ss:
                    b_ss = b.get("screenshot", "")
                if b_ss:
                    pname = Path(b_ss).name.lower()
                    ss_map[pname] = {
                        "type": f"Temuan Bug [{b.get('id', '')}]",
                        "title": b.get("title", ""),
                        "status": b.get("severity", "CRITICAL"),
                        "details": b.get("actual") or b.get("location") or ""
                    }

            add_h1("Bukti Visual Pengujian (Visual Evidence)")
            for img_idx, img_path in enumerate(screenshots, 1):
                try:
                    doc.add_picture(str(img_path), width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    info = ss_map.get(img_path.name.lower(), {})
                    context_title = info.get("title", img_path.stem.replace("-", " ").replace("_", " ").title())
                    context_type = info.get("type", "Dokumentasi")
                    context_status = info.get("status", "")
                    context_details = info.get("details", "")

                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r_cap = cap.add_run(f'Gambar {img_idx}: {context_title}')
                    r_cap.font.name = 'Calibri'
                    r_cap.font.size = Pt(10)
                    r_cap.bold = True
                    r_cap.font.color.rgb = RGBColor(0x1B, 0x2A, 0x8A)

                    if context_status or context_details:
                        cap_sub = doc.add_paragraph()
                        cap_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        status_str = f"[{context_status.upper()}] " if context_status else ""
                        r_sub = cap_sub.add_run(f'{status_str}({context_type}) {context_details}')
                        r_sub.font.name = 'Calibri'
                        r_sub.font.size = Pt(9)
                        r_sub.italic = True
                        r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    doc.add_paragraph('')
                except Exception as e:
                    print(f"Warning: Could not embed image: {e}")

        # Save document with lock retry
        filepath = self.output_dir / filename
        try:
            doc.save(str(filepath))
            print(f"DOCX report saved ({mode.upper()} layout): {filepath}")
            return filepath
        except PermissionError:
            alt_path = self.output_dir / f"QA_TEST_REPORT_{int(datetime.now().timestamp())}.docx"
            doc.save(str(alt_path))
            print(f"Warning: File {filepath} is locked (currently opened in Word). Saved as: {alt_path}")
            return alt_path

    # ── PREMIUM HTML ─────────────────────────────────────────
    def generate_html(self, filename="QA_TEST_REPORT.html"):
        score = self._quality_score()
        summary = self._summary()
        bugs = self._bugs()
        tests = self._tests()
        perf = self._perf()
        a11y = self._a11y()

        total = summary.get("total_tests", summary.get("total", 0))
        if not total and tests:
            total = len(tests)
            passed = sum(1 for t in tests if str(t.get("status", "")).strip().lower() in ["pass", "passed", "success", "ok"])
            failed = sum(1 for t in tests if str(t.get("status", "")).strip().lower() in ["fail", "failed", "error", "defect"])
            flaky_count = sum(1 for t in tests if t.get("flaky"))
        else:
            passed = summary.get("passed", 0)
            failed = summary.get("failed", 0)
            flaky_count = summary.get("flaky", 0)
        duration = summary.get("duration_ms", 0)
        duration_display = format_duration(duration)

        score_class = "good" if score >= 80 else "warn" if score >= 60 else "bad"
        readiness = "READY" if score >= 80 and not any(b.get("severity") == "CRITICAL" for b in bugs) else "NOT READY"
        readiness_class = "ready" if readiness == "READY" else "not-ready"

        # Video section
        video_html = ""
        video_path = self._find_video()
        if video_path:
            b64_video = self._embed_file(video_path, "video")
            p = Path(video_path)
            ext = p.suffix.lower().lstrip(".")
            mime = {"mp4": "video/mp4", "webm": "video/webm"}.get(ext, "video/webm")
            if b64_video:
                video_html = f'''
            <section class="report-section">
                <h2>Test Recording</h2>
                <p class="section-desc">Video rekaman proses testing otomatis</p>
                <div class="video-container">
                    <video controls preload="metadata" style="width:100%;max-height:500px;border-radius:8px;">
                        <source src="{b64_video}" type="{mime}">
                        Browser tidak support video tag.
                    </video>
                </div>
            </section>'''
            else:
                # fallback: link to file
                video_html = f'''
            <section class="report-section">
                <h2>Test Recording</h2>
                <p class="section-desc">Video rekaman proses testing otomatis</p>
                <div class="video-container">
                    <video controls preload="metadata" style="width:100%;max-height:500px;border-radius:8px;">
                        <source src="{video_path}" type="{mime}">
                        Browser tidak support video tag.
                    </video>
                </div>
            </section>'''

        # Bug rows
        bug_html = ""
        for bug in bugs:
            sev = bug.get("severity", "MEDIUM")
            priority = bug.get("priority", "P2")
            bug_type = bug.get("type", "Functional")
            steps = bug.get("steps", [])
            steps_html = "".join(f"<li>{strip_step_number(s)}</li>" for s in steps) if steps else "<li>N/A</li>"
            
            # Evidence screenshot extraction - handle string, dict, or direct field
            evidence = bug.get("evidence", "")
            screenshot = None
            video = None
            if isinstance(evidence, str) and evidence:
                screenshot = evidence
            elif isinstance(evidence, dict):
                screenshot = evidence.get("screenshot") or evidence.get("image")
                video = evidence.get("video")
            if not screenshot:
                screenshot = bug.get("screenshot") or bug.get("image")

            media_html = ""
            # Screenshot
            if screenshot:
                b64_img = self._embed_file(screenshot)
                img_src = b64_img if b64_img else screenshot
                media_html += f'<div class="evidence"><img src="{img_src}" alt="Bukti Bug" class="lightbox-trigger" data-caption="Bukti {bug.get("id","")}: {bug.get("title","")}"></div>'

            # Video per-bug (if exists)
            if video:
                b64_vid = self._embed_file(video, "video")
                if b64_vid:
                    media_html += f'''<div class="evidence">
                        <video controls preload="metadata" style="width:100%;max-height:300px;border-radius:8px;margin-top:8px;">
                            <source src="{b64_vid}" type="video/webm">
                        </video>
                    </div>'''

            bug_html += f'''
            <div class="bug-card {sev.lower()}">
                <div class="bug-header">
                    <span class="severity-badge {sev.lower()}">{sev}</span>
                    <span class="bug-id">{bug.get("id","")}</span>
                    <span class="bug-title">{bug.get("title","")}</span>
                    <span class="layer-badge" style="margin-left:auto; font-size:0.75rem;">Priority: {priority} | {bug_type}</span>
                </div>
                <div class="bug-body">
                    <div class="bug-field"><strong>Lokasi / URL:</strong> <code>{bug.get("location","")}</code></div>
                    <div class="bug-field"><strong>Langkah Reproduksi (Steps):</strong><ol>{steps_html}</ol></div>
                    <div class="bug-field"><strong>Kondisi Aktual (Actual):</strong> <span style="color:var(--danger)">{bug.get("actual","")}</span></div>
                    <div class="bug-field"><strong>Kondisi Harapan (Expected):</strong> <span style="color:var(--success)">{bug.get("expected","")}</span></div>
                    <div class="bug-field"><strong>Rekomendasi Solusi:</strong> <span style="color:var(--accent); font-weight:500;">{bug.get("recommendation","")}</span></div>
                    {media_html}
                </div>
            </div>'''

        # Test results grouped by Category or Endpoint Matrix (if API mode)
        test_results_html = ""
        mode = self.data.get("testing_mode", "e2e").lower()

        if mode == "api":
            # Specialized API Matrix View
            api_rows = ""
            for idx, t in enumerate(tests, 1):
                status_raw = str(t.get("status", "")).strip().lower()
                is_pass = status_raw in ["pass", "passed", "success", "ok"]
                status_cls = "pass" if is_pass else "fail"
                status_text = "200 OK / PASS" if is_pass else "ERROR / FAIL"
                dur_ms = t.get("duration_ms", t.get("duration", 0))
                
                # Payload / Request details
                req_method = t.get("method", "POST" if "create" in t.get("name","").lower() or "submit" in t.get("name","").lower() else "GET")
                endpoint_url = t.get("endpoint", t.get("location", self.data.get("target_url", "")))
                
                api_rows += f'''
                <tr>
                    <td style="text-align:center; font-weight:600;">{idx}</td>
                    <td>
                        <div style="display:flex; align-items:center; gap:6px; margin-bottom:2px;">
                            <span class="layer-badge" style="font-weight:700; font-size:0.75rem;">{req_method}</span>
                            <strong style="color:var(--text);">{t.get("name","")}</strong>
                        </div>
                        <div style="font-size:0.8rem; color:var(--text-secondary); font-family:monospace;">{endpoint_url}</div>
                    </td>
                    <td>
                        <div style="font-size:0.85rem;"><strong>Expected:</strong> <span style="color:var(--success)">{t.get("expected", "HTTP 200 & Valid Schema")}</span></div>
                        <div style="font-size:0.85rem; margin-top:2px;"><strong>Actual:</strong> <span style="color:{'var(--danger)' if not is_pass else 'var(--text)'}">{t.get("actual", t.get("details", "-"))}</span></div>
                    </td>
                    <td style="text-align:center; font-weight:600; font-size:0.85rem;">{format_duration(dur_ms)}</td>
                    <td style="text-align:center;"><span class="status-badge {status_cls}">{status_text}</span></td>
                </tr>'''

            test_results_html = f'''
            <div class="category-header"><span>Matriks Pengujian Endpoint API & Layanan</span><span class="category-count">{len(tests)} Endpoint</span></div>
            <table class="data-table">
                <thead>
                    <tr>
                        <th style="width:40px; text-align:center;">#</th>
                        <th>Endpoint & Nama Pengujian</th>
                        <th>Validasi Payload & Response</th>
                        <th style="width:110px; text-align:center;">Latensi</th>
                        <th style="width:130px; text-align:center;">Status</th>
                    </tr>
                </thead>
                <tbody>
                    {api_rows}
                </tbody>
            </table>'''
        else:
            # E2E & Bug Hunter Category Accordion View
            happy_tests = [t for t in tests if t.get("category") == "happy_path" or (not t.get("category") and str(t.get("status","")).strip().lower() in ["pass", "passed", "success", "ok"] and "invalid" not in t.get("name","").lower() and "fail" not in t.get("name","").lower())]
            negative_tests = [t for t in tests if t not in happy_tests]

            def _render_html_test_group(test_list, group_name):
                if not test_list:
                    return ""
                group_html = f'<div class="category-header"><span>{group_name}</span><span class="category-count">{len(test_list)} Skenario</span></div>'
                for idx, t in enumerate(test_list, 1):
                    status_raw = str(t.get("status", "")).strip().lower()
                    is_pass = status_raw in ["pass", "passed", "success", "ok"]
                    status_cls = "pass" if is_pass else "fail"
                    flaky_mark = '<span class="flaky-badge">FLAKY</span>' if t.get("flaky") else ""
                    dur_ms = t.get("duration_ms", t.get("duration", 0))
                    status_display = "PASSED" if is_pass else "FAILED"
                    
                    # Steps formatting
                    steps = t.get("steps", [])
                    if isinstance(steps, list) and steps:
                        steps_li = "".join([f"<li>{strip_step_number(s)}</li>" for s in steps])
                        steps_html = f"<ol>{steps_li}</ol>"
                    elif isinstance(steps, str) and steps:
                        steps_html = f"<p>{steps}</p>"
                    else:
                        steps_html = "<p style='color:var(--text-secondary)'>-</p>"

                    expected = t.get("expected", "-")
                    actual = t.get("actual", t.get("details", t.get("error", "-")))

                    # Screenshot button
                    evidence = t.get("evidence", {})
                    test_ss = t.get("screenshot") or (evidence.get("screenshot") if isinstance(evidence, dict) else None)
                    ss_html = ""
                    if test_ss:
                        b64 = self._embed_file(test_ss)
                        ss_src = b64 if b64 else test_ss
                        test_name_clean = t.get("name","").replace('"', '&quot;')
                        ss_html = f'<div style="margin-top:10px;"><button class="ss-link ss-btn" data-src="{ss_src}" data-caption="Bukti Skenario: {test_name_clean}">Lihat Screenshot Bukti</button></div>'

                    group_html += f'''
                    <div class="test-card">
                        <div class="test-card-header" onclick="this.nextElementSibling.style.display = this.nextElementSibling.style.display === 'none' ? 'block' : 'none'">
                            <div class="test-card-title">
                                <span class="status-badge {status_cls}">{status_display}</span>
                                <span>#{idx} {t.get("name","")}</span>
                                {flaky_mark}
                            </div>
                            <div class="test-card-meta">
                                <span class="layer-badge">{t.get("layer","UI").upper()}</span>
                                <span style="font-size:0.8rem;color:var(--text-secondary)">{format_duration(dur_ms)}</span>
                            </div>
                        </div>
                        <div class="test-card-body" style="display:{'block' if not is_pass else 'none'};">
                            <div class="test-field"><strong>Steps / Langkah:</strong> {steps_html}</div>
                            <div class="test-field"><strong>Expected:</strong> <span style="color:var(--success)">{expected}</span></div>
                            <div class="test-field"><strong>Actual:</strong> <span style="color:{'var(--danger)' if not is_pass else 'var(--text)'}">{actual}</span></div>
                            <div class="test-field"><strong>Browser:</strong> {t.get("browser","Chromium")} | <strong>Retries:</strong> {t.get("retries", 0)}</div>
                            {ss_html}
                        </div>
                    </div>'''
                return group_html

            if happy_tests:
                test_results_html += _render_html_test_group(happy_tests, "Skenario Happy Path (Positif)")
            if negative_tests:
                test_results_html += _render_html_test_group(negative_tests, "Skenario Negative / Boundary / Exception")
            if not tests:
                test_results_html = '<p style="color:var(--text-secondary);text-align:center;padding:20px;">Tidak ada test results tercatat.</p>'

        # Performance rows with full names + descriptions
        perf_html = ""
        if perf:
            perf_rows = ""
            for metric, val in perf.items():
                metric_clean = metric.replace("_ms", "").replace("_", " ").strip().upper()
                key_map = {
                    "PAGE LOAD MS": "PAGE LOAD",
                    "PAGE LOAD": "PAGE LOAD",
                    "TTFB MS": "TTFB",
                    "FCP MS": "FCP",
                    "LCP MS": "LCP",
                    "CLS": "CLS",
                    "FID": "FID",
                    "INP": "INP",
                    "TBT": "TBT"
                }
                display_key = key_map.get(metric_clean, metric_clean)
                info = METRIC_INFO.get(display_key, METRIC_INFO.get(metric_clean, {}))
                full_name = info.get("full", display_key)
                desc = info.get("desc", "")
                if isinstance(val, dict):
                    v = val.get("value", "")
                    u = val.get("unit", "")
                    r = val.get("rating", "unknown")
                else:
                    v = val
                    u = ""
                    r = "unknown"
                bar_pct = min(100, max(5, 100 - int(float(v) / 50) if r != "unknown" else 50))
                perf_rows += f'''
                <tr>
                    <td>
                        <strong style="color:var(--text); font-size:0.95rem;">{display_key}</strong>
                        <div class="metric-fullname" style="font-weight:600; color:var(--primary); margin-top:2px;">{full_name}</div>
                        <div class="metric-desc" style="margin-top:4px; line-height:1.4; color:var(--text-secondary);">{desc}</div>
                    </td>
                    <td style="font-weight:700; font-size:1rem; white-space:nowrap;">{v}{u}</td>
                    <td><div class="perf-bar"><div class="perf-fill {r}" style="width:{bar_pct}%"></div></div></td>
                    <td><span class="rating-badge {r}">{r.replace("_"," ").title()}</span></td>
                </tr>'''
            perf_html = f'''
            <section class="report-section">
                <h2>Performance (Core Web Vitals & Kecepatan Akses)</h2>
                <p class="section-desc">Pengukuran metrik kecepatan muat, responsivitas, dan stabilitas visual sesuai standar industri Google Web Vitals.</p>
                <table class="data-table">
                    <thead><tr><th style="width:50%;">Metrik & Deskripsi Penjelasan</th><th style="width:15%;">Nilai Uji</th><th style="width:20%;">Indikator</th><th style="width:15%;">Rating</th></tr></thead>
                    <tbody>{perf_rows}</tbody>
                </table>
            </section>'''

        # Accessibility rows
        a11y_html = ""
        if a11y:
            a11y_rows = ""
            for issue in a11y:
                sev = issue.get("severity", "moderate").lower()
                desc = issue.get("description", "")
                rule = issue.get("type", issue.get("rule", "Aksesibilitas UI"))
                wcag_raw = issue.get("wcag", "")

                # Friendly WCAG explanations
                wcag_map = {
                    "1.1.1": "1.1.1 Non-text Content (Teks Alternatif Gambar)",
                    "1.3.1": "1.3.1 Info & Relationships (Struktur Form & Label)",
                    "1.4.3": "1.4.3 Contrast Minimum (Kontras Warna Teks)",
                    "2.1.1": "2.1.1 Keyboard Accessible (Akses Tombol/Keyboard)",
                    "2.4.7": "2.4.7 Focus Visible (Indikator Fokus Terlihat)",
                    "3.3.2": "3.3.2 Labels or Instructions (Petunjuk Pengisian)"
                }
                wcag_friendly = wcag_raw
                for k, v in wcag_map.items():
                    if k in wcag_raw:
                        wcag_friendly = f"WCAG 2.1 AA — {v}"
                        break
                if not wcag_friendly or wcag_friendly == wcag_raw:
                    wcag_friendly = f"WCAG 2.1 AA — {wcag_raw}" if wcag_raw else "WCAG 2.1 AA (Standar Aksesibilitas)"

                # Friendly description elaboration
                friendly_desc = desc
                if "lacks explicit aria-label" in desc.lower() or "form label missing" in rule.lower():
                    friendly_desc = f"<strong>Input form belum memiliki label/teks pembaca layar:</strong> {desc}"
                elif "contrast ratio is below" in desc.lower() or "color contrast" in rule.lower():
                    friendly_desc = f"<strong>Kontras warna teks terlalu redup sehingga sulit dibaca pengguna:</strong> {desc}"

                a11y_rows += f'''
                <tr>
                    <td style="white-space:nowrap;"><span class="severity-badge {sev}">{sev.upper()}</span></td>
                    <td>
                        <div style="font-weight:600; color:var(--text); margin-bottom:3px;">{rule}</div>
                        <div style="font-size:0.85rem; color:var(--text-secondary); line-height:1.4;">{friendly_desc}</div>
                    </td>
                    <td><code style="font-size:0.8rem; background:var(--bg); padding:3px 6px; border-radius:4px; border:1px solid var(--border);">{issue.get("element","")}</code></td>
                    <td>
                        <div style="font-size:0.85rem; font-weight:600; color:var(--accent);">{wcag_friendly}</div>
                    </td>
                </tr>'''
            a11y_html = f'''
            <section class="report-section">
                <h2>Uji Aksesibilitas & Kenyamanan Pengguna (WCAG 2.1 AA)</h2>
                <p class="section-desc">Pemeriksaan kepatuhan antarmuka terhadap standar aksesibilitas internasional (kemudahan membaca teks, kelengkapan label form, dan kenyamanan navigasi bagi semua pengguna).</p>
                <table class="data-table">
                    <thead><tr><th style="width:100px;">Dampak</th><th>Jenis Masalah & Penjelasan</th><th>Elemen UI Terkait</th><th>Standar Patokan</th></tr></thead>
                    <tbody>{a11y_rows}</tbody>
                </table>
            </section>'''

        # Visual Evidence Gallery (All screenshots)
        gallery_html = ""
        ss_dir = self.output_dir / "screenshots"
        screenshots = []
        if ss_dir.exists():
            for ext in ["*.png", "*.jpg", "*.jpeg", "*.webp"]:
                screenshots.extend(sorted(ss_dir.glob(ext)))

        if screenshots:
            # Map screenshots to tests / bugs for richer descriptions
            ss_map = {}
            for t in tests:
                t_ss = t.get("screenshot") or ""
                if isinstance(t.get("evidence"), dict):
                    t_ss = t_ss or t["evidence"].get("screenshot", "")
                if t_ss:
                    pname = Path(t_ss).name.lower()
                    ss_map[pname] = {
                        "type": "Test Skenario",
                        "title": t.get("name", ""),
                        "status": t.get("status", "PASS"),
                        "details": t.get("actual") or t.get("details") or t.get("expected") or ""
                    }
            for b in bugs:
                b_ev = b.get("evidence") or {}
                b_ss = b_ev.get("screenshot") if isinstance(b_ev, dict) else ""
                if not b_ss:
                    b_ss = b.get("screenshot", "")
                if b_ss:
                    pname = Path(b_ss).name.lower()
                    ss_map[pname] = {
                        "type": f"Temuan Bug [{b.get('id', '')}]",
                        "title": b.get("title", ""),
                        "status": b.get("severity", "CRITICAL"),
                        "details": b.get("actual") or b.get("location") or ""
                    }

            gallery_items = ""
            for img_idx, img_path in enumerate(screenshots, 1):
                b64_img = self._embed_file(str(img_path))
                img_name = img_path.stem.replace("-", " ").replace("_", " ").title()
                src = b64_img if b64_img else str(img_path)
                
                # Context info
                info = ss_map.get(img_path.name.lower(), {})
                context_type = info.get("type", "Dokumentasi Pengujian")
                context_title = info.get("title", img_name)
                context_status = info.get("status", "")
                context_details = info.get("details", "")

                status_badge_html = ""
                if context_status:
                    st_low = context_status.lower()
                    if st_low in ["pass", "passed", "success", "ok"]:
                        status_badge_html = f'<span class="status-badge pass" style="font-size:0.75rem; padding:2px 8px;">{context_status.upper()}</span>'
                    elif st_low in ["fail", "failed", "error", "critical", "high"]:
                        status_badge_html = f'<span class="status-badge fail" style="font-size:0.75rem; padding:2px 8px;">{context_status.upper()}</span>'
                    else:
                        status_badge_html = f'<span class="status-badge" style="background:#e5e7eb; font-size:0.75rem; padding:2px 8px;">{context_status.upper()}</span>'

                details_html = f'<div style="font-size:0.8rem; color:var(--text-secondary); margin-top:4px; line-height:1.4;">{context_details}</div>' if context_details else ''

                gallery_items += f'''
                <div class="gallery-card">
                    <div class="gallery-img-wrap">
                        <img src="{src}" alt="{img_name}" loading="lazy" class="lightbox-trigger" data-caption="Gambar {img_idx}: {context_title}">
                    </div>
                    <div class="gallery-caption">
                        <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:4px; gap:8px;">
                            <strong style="color:var(--accent); font-size:0.85rem;">Gambar {img_idx}</strong>
                            {status_badge_html}
                        </div>
                        <div style="font-weight:600; font-size:0.9rem; margin-bottom:2px; color:var(--text);">{context_title}</div>
                        <div style="font-size:0.75rem; color:var(--text-secondary); text-transform:uppercase; letter-spacing:0.5px; font-weight:600;">{context_type}</div>
                        {details_html}
                    </div>
                </div>'''
            
            gallery_html = f'''
            <section class="report-section">
                <h2>Bukti Visual Pengujian (Visual Evidence)</h2>
                <p class="section-desc">Koleksi tangkapan layar bukti eksekusi pengujian, form UI, hasil validasi, dan temuan kendala sistem beserta penjelasannya.</p>
                <div class="gallery-grid">
                    {gallery_items}
                </div>
            </section>'''

        # Severity breakdown
        crit = sum(1 for b in bugs if b.get("severity") == "CRITICAL")
        high = sum(1 for b in bugs if b.get("severity") == "HIGH")
        med = sum(1 for b in bugs if b.get("severity") == "MEDIUM")
        low = sum(1 for b in bugs if b.get("severity") == "LOW")

        # Build HTML — note: we use string concat instead of f-string for the JS
        # to avoid {{ }} escaping issues with the toggle function
        html_parts = []
        html_parts.append('''<!DOCTYPE html>
<html lang="en" data-theme="light">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>QA Test Report</title>
    <style>
        :root {
            --bg: #f0f2f5; --surface: #ffffff; --text: #1a1a2e; --text-secondary: #6b7280;
            --border: #e5e7eb; --accent: #3b82f6; --success: #10b981; --danger: #ef4444;
            --warning: #f59e0b; --orange: #f97316; --info: #6366f1;
        }
        [data-theme="dark"] {
            --bg: #0f172a; --surface: #1e293b; --text: #e2e8f0; --text-secondary: #94a3b8;
            --border: #334155; --accent: #60a5fa;
        }
        * { margin:0; padding:0; box-sizing:border-box; }
        body { font-family: 'Segoe UI', system-ui, -apple-system, sans-serif; background:var(--bg); color:var(--text); line-height:1.6; }
        .container { max-width:1280px; margin:0 auto; padding:24px; }

        /* Header */
        .report-header { background: linear-gradient(135deg, #1e293b 0%, #334155 100%); color:#fff; padding:40px; border-radius:16px; margin-bottom:24px; position:relative; overflow:hidden; }
        .report-header::after { content:''; position:absolute; top:-50%; right:-20%; width:400px; height:400px; background:radial-gradient(circle, rgba(59,130,246,0.15) 0%, transparent 70%); }
        .report-header h1 { font-size:2rem; font-weight:700; margin-bottom:4px; }
        .report-header .meta { opacity:0.8; font-size:0.9rem; }
        .report-header .meta span { margin-right:20px; }
        .theme-toggle { position:absolute; top:20px; right:20px; background:rgba(255,255,255,0.15); border:none; color:#fff; padding:8px 14px; border-radius:8px; cursor:pointer; font-size:0.85rem; backdrop-filter:blur(4px); transition: background 0.2s; }
        .theme-toggle:hover { background:rgba(255,255,255,0.25); }
        .readiness { display:inline-block; margin-top:12px; padding:6px 18px; border-radius:20px; font-weight:700; font-size:0.95rem; }
        .readiness.ready { background:#10b981; color:#fff; }
        .readiness.not-ready { background:#ef4444; color:#fff; }

        /* Summary Cards */
        .summary-grid { display:grid; grid-template-columns: repeat(auto-fit, minmax(180px,1fr)); gap:16px; margin-bottom:24px; }
        .summary-card { background:var(--surface); border:1px solid var(--border); border-radius:12px; padding:20px; text-align:center; transition:transform 0.2s, box-shadow 0.2s; }
        .summary-card:hover { transform:translateY(-2px); box-shadow:0 8px 25px rgba(0,0,0,0.08); }
        .summary-card .value { font-size:2.2rem; font-weight:800; line-height:1; }
        .summary-card .label { font-size:0.85rem; color:var(--text-secondary); margin-top:4px; }
        .summary-card .value.good { color:var(--success); }
        .summary-card .value.warn { color:var(--warning); }
        .summary-card .value.bad { color:var(--danger); }
        .summary-card .value.pass { color:var(--success); }
        .summary-card .value.fail { color:var(--danger); }
        .summary-card .value.flaky { color:var(--warning); }

        /* Sections */
        .report-section { background:var(--surface); border:1px solid var(--border); border-radius:12px; padding:24px; margin-bottom:24px; }
        .report-section h2 { font-size:1.3rem; margin-bottom:6px; padding-bottom:8px; border-bottom:2px solid var(--border); }
        .section-desc { font-size:0.85rem; color:var(--text-secondary); margin-bottom:16px; }

        /* Video */
        .video-container { border-radius:8px; overflow:hidden; background:#000; }

        /* Severity Breakdown */
        .severity-breakdown { display:flex; gap:12px; flex-wrap:wrap; margin-bottom:20px; }
        .sev-chip { display:flex; align-items:center; gap:6px; padding:6px 14px; border-radius:8px; font-size:0.85rem; font-weight:600; }
        .sev-chip.critical { background:#fef2f2; color:#dc2626; border:1px solid #fecaca; }
        .sev-chip.high { background:#fff7ed; color:#ea580c; border:1px solid #fed7aa; }
        .sev-chip.medium { background:#fffbeb; color:#d97706; border:1px solid #fde68a; }
        .sev-chip.low { background:#f0fdf4; color:#16a34a; border:1px solid #bbf7d0; }
        [data-theme="dark"] .sev-chip.critical { background:#450a0a; border-color:#991b1b; }
        [data-theme="dark"] .sev-chip.high { background:#431407; border-color:#9a3412; }
        [data-theme="dark"] .sev-chip.medium { background:#451a03; border-color:#92400e; }
        [data-theme="dark"] .sev-chip.low { background:#052e16; border-color:#166534; }

        /* Bug Cards */
        .bug-card { border:1px solid var(--border); border-radius:10px; margin-bottom:16px; overflow:hidden; transition:box-shadow 0.2s; }
        .bug-card:hover { box-shadow:0 4px 15px rgba(0,0,0,0.06); }
        .bug-card.critical { border-left:4px solid #ef4444; }
        .bug-card.high { border-left:4px solid #f97316; }
        .bug-card.medium { border-left:4px solid #f59e0b; }
        .bug-card.low { border-left:4px solid #10b981; }
        .bug-header { padding:14px 18px; background:var(--bg); display:flex; align-items:center; gap:10px; flex-wrap:wrap; }
        .bug-body { padding:18px; }
        .bug-field { margin-bottom:10px; }
        .bug-field ol { margin-left:20px; margin-top:4px; }
        .bug-field ol li { margin-bottom:2px; }
        .evidence { margin-top:12px; }
        .evidence img { max-width:100%; border-radius:8px; border:1px solid var(--border); box-shadow:0 2px 8px rgba(0,0,0,0.06); }

        /* Visual Gallery */
        .gallery-grid { display:grid; grid-template-columns: repeat(auto-fit, minmax(280px, 1fr)); gap:20px; margin-top:12px; }
        .gallery-card { background:var(--bg); border:1px solid var(--border); border-radius:10px; overflow:hidden; transition:transform 0.2s, box-shadow 0.2s; }
        .gallery-card:hover { transform:translateY(-2px); box-shadow:0 6px 20px rgba(0,0,0,0.1); }
        .gallery-img-wrap { overflow:hidden; background:#000; display:flex; align-items:center; justify-content:center; max-height:240px; }
        .gallery-img-wrap img { width:100%; height:auto; object-fit:contain; transition:transform 0.3s ease; }
        .gallery-img-wrap:hover img { transform:scale(1.03); }
        .gallery-caption { padding:12px 14px; font-size:0.85rem; color:var(--text); border-top:1px solid var(--border); }

        /* Badges */
        .severity-badge { display:inline-block; padding:3px 10px; border-radius:6px; font-size:0.75rem; font-weight:700; color:#fff; text-transform:uppercase; }
        .severity-badge.critical { background:#ef4444; }
        .severity-badge.high { background:#f97316; }
        .severity-badge.medium { background:#f59e0b; color:#1a1a2e; }
        .severity-badge.low { background:#10b981; }
        .severity-badge.serious { background:#f97316; }
        .severity-badge.moderate { background:#f59e0b; color:#1a1a2e; }
        .severity-badge.minor { background:#10b981; }
        .severity-badge.violation { background:#ef4444; }
        .status-badge { padding:3px 10px; border-radius:6px; font-size:0.8rem; font-weight:600; }
        .status-badge.pass { background:#d1fae5; color:#065f46; }
        .status-badge.fail { background:#fee2e2; color:#991b1b; }
        [data-theme="dark"] .status-badge.pass { background:#064e3b; color:#6ee7b7; }
        [data-theme="dark"] .status-badge.fail { background:#450a0a; color:#fca5a5; }
        .flaky-badge { background:#fef3c7; color:#92400e; padding:2px 8px; border-radius:6px; font-size:0.7rem; font-weight:600; margin-left:6px; }
        .layer-badge { background:var(--bg); padding:2px 8px; border-radius:6px; font-size:0.8rem; font-weight:500; }
        .rating-badge { padding:3px 10px; border-radius:6px; font-size:0.8rem; font-weight:600; }
        .rating-badge.good { background:#d1fae5; color:#065f46; }
        .rating-badge.needs_improvement { background:#fef3c7; color:#92400e; }
        .rating-badge.poor { background:#fee2e2; color:#991b1b; }
        [data-theme="dark"] .rating-badge.good { background:#064e3b; color:#6ee7b7; }
        [data-theme="dark"] .rating-badge.needs_improvement { background:#451a03; color:#fbbf24; }
        [data-theme="dark"] .rating-badge.poor { background:#450a0a; color:#fca5a5; }
        .bug-id { font-weight:700; font-family:monospace; }
        .bug-title { font-weight:600; }

        /* Test Accordion / Details */
        .test-card { border:1px solid var(--border); border-radius:8px; margin-bottom:12px; overflow:hidden; background:var(--surface); }
        .test-card-header { padding:12px 16px; display:flex; align-items:center; justify-content:space-between; gap:12px; cursor:pointer; background:var(--bg); }
        .test-card-header:hover { background:rgba(0,0,0,0.02); }
        [data-theme="dark"] .test-card-header:hover { background:rgba(255,255,255,0.02); }
        .test-card-title { display:flex; align-items:center; gap:8px; font-weight:600; }
        .test-card-meta { display:flex; align-items:center; gap:8px; }
        .test-card-body { padding:14px 18px; border-top:1px solid var(--border); font-size:0.9rem; }
        .test-field { margin-bottom:8px; }
        .test-field strong { color:var(--text); }
        .test-field ol { margin-left:20px; margin-top:4px; }
        .test-field ol li { margin-bottom:2px; }
        .category-header { font-size:1.1rem; font-weight:700; color:var(--text); margin:20px 0 10px 0; border-bottom:2px solid var(--border); padding-bottom:6px; display:flex; align-items:center; justify-content:space-between; }
        .category-count { font-size:0.85rem; font-weight:500; color:var(--text-secondary); }

        /* Lightbox Modal */
        .lightbox-trigger { cursor: zoom-in; }
        .lightbox-modal { display:none; position:fixed; z-index:9999; inset:0; background:rgba(0,0,0,0.88); backdrop-filter:blur(6px); justify-content:center; align-items:center; flex-direction:column; padding:20px; }
        .lightbox-modal.active { display:flex; animation:fadeIn 0.2s ease; }
        .lightbox-content { max-width:92vw; max-height:85vh; border-radius:8px; box-shadow:0 10px 40px rgba(0,0,0,0.5); object-fit:contain; }
        .lightbox-caption { color:#e2e8f0; font-size:0.95rem; margin-top:12px; text-align:center; max-width:800px; }
        .lightbox-close { position:absolute; top:20px; right:24px; color:#fff; font-size:2rem; font-weight:bold; cursor:pointer; background:rgba(255,255,255,0.15); width:44px; height:44px; border-radius:50%; display:flex; align-items:center; justify-content:center; border:none; transition:background 0.2s; }
        .lightbox-close:hover { background:rgba(255,255,255,0.3); }
        @keyframes fadeIn { from { opacity:0; } to { opacity:1; } }

        /* Metric descriptions */
        .metric-fullname { font-size:0.8rem; color:var(--text-secondary); font-weight:500; }
        .metric-desc { font-size:0.75rem; color:var(--text-secondary); font-style:italic; margin-top:2px; }

        /* Tables */
        .data-table { width:100%; border-collapse:collapse; }
        .data-table th { text-align:left; padding:10px 14px; background:var(--bg); font-size:0.85rem; font-weight:600; color:var(--text-secondary); border-bottom:2px solid var(--border); }
        .data-table td { padding:10px 14px; border-bottom:1px solid var(--border); font-size:0.9rem; }
        .data-table tr:hover td { background:rgba(59,130,246,0.04); }

        /* Perf bars */
        .perf-bar { width:100%; height:8px; background:var(--bg); border-radius:4px; overflow:hidden; }
        .perf-fill { height:100%; border-radius:4px; transition:width 0.5s ease; }
        .perf-fill.good { background:var(--success); }
        .perf-fill.needs_improvement { background:var(--warning); }
        .perf-fill.poor { background:var(--danger); }

        /* Footer */
        .report-footer { text-align:center; padding:20px; color:var(--text-secondary); font-size:0.8rem; }

        /* Print */
        @media print {
            body { background:#fff; }
            .theme-toggle { display:none; }
            .report-header { background:#1e293b !important; -webkit-print-color-adjust:exact; print-color-adjust:exact; }
            .summary-card:hover { transform:none; box-shadow:none; }
            .video-container { display:none; }
        }
        @media (max-width: 768px) {
            .container { padding:12px; }
            .report-header { padding:24px; }
            .report-header h1 { font-size:1.4rem; }
            .summary-grid { grid-template-columns: repeat(2, 1fr); }
        }
    </style>
</head>
<body>
    <div class="container">
''')

        # Header section with dynamic data
        mode = self.data.get("testing_mode", "e2e").lower()
        mode_titles = {
            "e2e": "E2E Comprehensive Test Report",
            "bug_hunter": "Bug Hunter & Defect Audit Report",
            "api": "API & Service Integration Test Report"
        }
        report_title = mode_titles.get(mode, "QA Test Report")
        mode_badge = mode.upper().replace("_", " ")

        html_parts.append(f'''        <header class="report-header">
            <button class="theme-toggle" id="themeBtn">Dark / Light</button>
            <div style="display:flex; align-items:center; gap:10px; margin-bottom:6px;">
                <span class="status-badge" style="background:rgba(255,255,255,0.2); color:#fff; font-size:0.75rem; border:1px solid rgba(255,255,255,0.3);">{mode_badge}</span>
            </div>
            <h1 style="margin-bottom:8px;">{report_title}</h1>
            <div class="meta">
                <span><strong>Project:</strong> {self.data.get("project_name","N/A")}</span>
                <span><strong>Target:</strong> {self.data.get("target_url","N/A")}</span>
                <span><strong>Generated:</strong> {self.data.get("generated_at", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))}</span>
            </div>
            <div class="readiness {readiness_class}">Production: {readiness}</div>
        </header>

        <div class="summary-grid">
            <div class="summary-card">
                <div class="value {score_class}">{score}</div>
                <div class="label">Quality Score</div>
            </div>
            <div class="summary-card">
                <div class="value">{total}</div>
                <div class="label">Total Tests</div>
            </div>
            <div class="summary-card">
                <div class="value pass">{passed}</div>
                <div class="label">Passed</div>
            </div>
            <div class="summary-card">
                <div class="value fail">{failed}</div>
                <div class="label">Failed</div>
            </div>
            <div class="summary-card">
                <div class="value flaky">{flaky_count}</div>
                <div class="label">Flaky</div>
            </div>
            <div class="summary-card">
                <div class="value">{duration_display}</div>
                <div class="label">Duration</div>
            </div>
        </div>

        {video_html}

        <section class="report-section">
            <h2>Bug Summary ({len(bugs)} found)</h2>
            <div class="severity-breakdown">
                <div class="sev-chip critical">CRITICAL: {crit}</div>
                <div class="sev-chip high">HIGH: {high}</div>
                <div class="sev-chip medium">MEDIUM: {med}</div>
                <div class="sev-chip low">LOW: {low}</div>
            </div>
            {bug_html if bug_html else '<p style="color:var(--text-secondary)">No bugs found.</p>'}
        </section>

        <section class="report-section">
            <h2>Skenario Pengujian (Happy Path & Negative Test Breakdown)</h2>
            <p class="section-desc">Rincian skenario pengujian beserta langkah (Steps), Expected, dan Actual Result.</p>
            {test_results_html}
        </section>

        {perf_html}
        {a11y_html}
        {gallery_html}

        <footer class="report-footer">
            QA Ultra Tester v2.1 &mdash; Generated from test-results.json (single source of truth)
        </footer>
    </div>

    <!-- Lightbox Modal -->
    <div id="lightboxModal" class="lightbox-modal">
        <button class="lightbox-close" id="lightboxClose">&times;</button>
        <img class="lightbox-content" id="lightboxImg" src="" alt="Enlarged View">
        <div class="lightbox-caption" id="lightboxCaption"></div>
    </div>
''')

        # JavaScript — separate to avoid f-string escaping issues
        html_parts.append('''    <script>
        (function() {
            var btn = document.getElementById('themeBtn');
            var html = document.documentElement;

            // Load saved theme
            var saved = localStorage.getItem('qa-report-theme');
            if (saved) {
                html.setAttribute('data-theme', saved);
            }

            btn.addEventListener('click', function() {
                var current = html.getAttribute('data-theme');
                var next = current === 'dark' ? 'light' : 'dark';
                html.setAttribute('data-theme', next);
                localStorage.setItem('qa-report-theme', next);
            });

            // Lightbox Modal Logic
            var modal = document.getElementById('lightboxModal');
            var modalImg = document.getElementById('lightboxImg');
            var modalCaption = document.getElementById('lightboxCaption');
            var closeBtn = document.getElementById('lightboxClose');

            function openLightbox(src, caption) {
                if (!src) return;
                modalImg.src = src;
                modalCaption.textContent = caption || '';
                modal.classList.add('active');
                document.body.style.overflow = 'hidden';
            }

            function closeLightbox() {
                modal.classList.remove('active');
                modalImg.src = '';
                document.body.style.overflow = '';
            }

            // Bind image clicks with .lightbox-trigger
            document.querySelectorAll('.lightbox-trigger').forEach(function(img) {
                img.addEventListener('click', function() {
                    openLightbox(this.src, this.getAttribute('data-caption') || this.alt);
                });
            });

            // Bind table SS buttons
            document.querySelectorAll('.ss-btn').forEach(function(btn) {
                btn.addEventListener('click', function(e) {
                    e.preventDefault();
                    var src = this.getAttribute('data-src');
                    var cap = this.getAttribute('data-caption');
                    openLightbox(src, cap);
                });
            });

            // Close actions
            closeBtn.addEventListener('click', closeLightbox);
            modal.addEventListener('click', function(e) {
                if (e.target === modal) {
                    closeLightbox();
                }
            });
            document.addEventListener('keydown', function(e) {
                if (e.key === 'Escape' && modal.classList.contains('active')) {
                    closeLightbox();
                }
            });
        })();
    </script>
</body>
</html>''')

        html = "".join(html_parts)

        filepath = self.output_dir / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html)
        print(f"HTML report saved: {filepath}")
        return filepath

    # ── JSON (write back / create) ───────────────────────────
    def generate_json(self, filename="test-results.json"):
        """Write current data back as JSON (useful for programmatic use)."""
        filepath = self.output_dir / filename
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(self.data, f, indent=2, ensure_ascii=False)
        print(f"JSON report saved: {filepath}")
        return filepath


# ── CLI entry point ──────────────────────────────────────────
if __name__ == "__main__":
    output_dir = sys.argv[1] if len(sys.argv) > 1 else "hasil-test"
    gen = QAReportGenerator(output_dir)

    if gen.load_from_json():
        gen.generate_html()
        gen.generate_docx()
        print(f"\nAll reports generated in {gen.output_dir}")
    else:
        # Demo mode with sample data
        print("No test-results.json found. Generating demo report...")
        demo_data = {
            "project_name": "Demo Project",
            "target_url": "https://example.com",
            "generated_at": datetime.now().isoformat(),
            "quality_score": 72,
            "summary": {"total": 15, "passed": 11, "failed": 3, "flaky": 1, "duration_ms": 185000},
            "testing_mode": "e2e",
            "bugs": [
                {
                    "id": "BUG-001", "severity": "CRITICAL",
                    "title": "SQL Injection on Login",
                    "location": "https://example.com/login",
                    "steps": ["Open login page", "Input SQL payload in email field", "Click login button"],
                    "actual": "Login successful without valid credentials",
                    "expected": "Login should fail with invalid credentials",
                    "evidence": {"screenshot": "", "video": "", "network_log": ""},
                    "recommendation": "Implement parameterized queries"
                },
                {
                    "id": "BUG-002", "severity": "HIGH",
                    "title": "XSS Reflected on Search",
                    "location": "https://example.com/search",
                    "steps": ["Navigate to search page", "Input <script>alert(1)</script> in search box", "Submit form"],
                    "actual": "Script executed in browser",
                    "expected": "Input should be sanitized",
                    "evidence": {"screenshot": "", "video": ""},
                    "recommendation": "Escape all user input in output"
                }
            ],
            "test_results": [
                {
                    "name": "Login Berhasil dengan Akun Valid",
                    "category": "happy_path",
                    "layer": "UI",
                    "status": "PASS",
                    "duration_ms": 1250,
                    "browser": "Chromium",
                    "steps": [
                        "1. Buka halaman login",
                        "2. Input username 'admin' dan password valid",
                        "3. Klik tombol Masuk"
                    ],
                    "expected": "Berhasil login dan diarahkan ke Dashboard utama",
                    "actual": "Login berhasil dan Dashboard tampil lengkap",
                    "retries": 0,
                    "flaky": False,
                    "details": ""
                },
                {
                    "name": "Input Data Form Banner Baru",
                    "category": "happy_path",
                    "layer": "UI",
                    "status": "PASS",
                    "duration_ms": 3400,
                    "browser": "Chromium",
                    "steps": [
                        "1. Buka menu Banner Input",
                        "2. Klik tombol Tambah Data",
                        "3. Isi Title, Upload Gambar banner.png, pilih Periode Tanggal",
                        "4. Klik tombol Simpan"
                    ],
                    "expected": "Data banner berhasil tersimpan dan muncul di daftar tabel",
                    "actual": "Data banner tersimpan dan tabel auto-refresh",
                    "retries": 0,
                    "flaky": False,
                    "details": ""
                },
                {
                    "name": "Submit Form Banner Tanpa Mengisi Field Wajib",
                    "category": "negative",
                    "layer": "UI",
                    "status": "FAIL",
                    "duration_ms": 2100,
                    "browser": "Chromium",
                    "steps": [
                        "1. Buka modal Tambah Data Banner",
                        "2. Kosongkan field Title dan File Gambar",
                        "3. Klik tombol Simpan langsung"
                    ],
                    "expected": "Muncul validasi error merah 'Field ini wajib diisi' dan form tidak tersubmit",
                    "actual": "Form tetap tersubmit dan server mengembalikan error 500",
                    "retries": 2,
                    "flaky": False,
                    "details": "Timeout waiting for validation message",
                    "screenshot": ""
                },
                {
                    "name": "API GET /api/v1/banners Endpoint Validation",
                    "category": "happy_path",
                    "layer": "API",
                    "status": "PASS",
                    "duration_ms": 450,
                    "browser": "Chromium",
                    "steps": [
                        "1. Kirim request HTTP GET /api/v1/banners dengan bearer token",
                        "2. Validasi response code dan schema data"
                    ],
                    "expected": "HTTP 200 OK dengan format JSON schema valid",
                    "actual": "HTTP 200 OK diterima dalam 450ms",
                    "retries": 0,
                    "flaky": False,
                    "details": ""
                },
                {
                    "name": "Upload File Gambar Ekstensi Tidak Didukung (.exe)",
                    "category": "negative",
                    "layer": "UI",
                    "status": "FAIL",
                    "duration_ms": 1800,
                    "browser": "Chromium",
                    "steps": [
                        "1. Buka form upload gambar banner",
                        "2. Pilih file payload.exe",
                        "3. Klik tombol Upload"
                    ],
                    "expected": "Ditolak oleh client-side validator: 'Hanya file JPG/PNG diperbolehkan'",
                    "actual": "File .exe lolos dan berhasil terupload ke server",
                    "retries": 1,
                    "flaky": False,
                    "details": "Client validator bypassed",
                    "screenshot": ""
                }
            ],
            "performance": {
                "LCP": {"value": 2300, "unit": "ms", "rating": "good"},
                "FID": {"value": 45, "unit": "ms", "rating": "good"},
                "CLS": {"value": 0.05, "unit": "", "rating": "good"},
                "TTFB": {"value": 820, "unit": "ms", "rating": "needs_improvement"},
                "FCP": {"value": 1200, "unit": "ms", "rating": "good"}
            },
            "accessibility": [
                {"type": "color_contrast", "element": ".text-muted", "description": "Contrast ratio 2.5:1 below 4.5:1 minimum", "wcag": "1.4.3", "severity": "serious"},
                {"type": "missing_alt", "element": "img.hero-banner", "description": "Image missing alt attribute", "wcag": "1.1.1", "severity": "critical"}
            ]
        }
        gen.set_data(demo_data)
        gen.generate_json()
        gen.generate_html()
        gen.generate_docx()
        print(f"\nDemo reports generated in {gen.output_dir}")
